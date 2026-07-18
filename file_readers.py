"""Extract raw text from resume files (PDF / DOCX)."""

from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class UnsupportedFileTypeError(ValueError):
    pass


def read_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    chunks = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            chunks.append(text)
    return "\n".join(chunks)


def read_docx(file_path: Path) -> str:
    document = Document(file_path)
    chunks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)
    return "\n".join(chunks)


def read_resume(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(file_path)
    if suffix == ".docx":
        return read_docx(file_path)
    raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")
